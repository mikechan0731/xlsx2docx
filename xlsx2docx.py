# -*-coding: utf-8-*-
# editor: MikeChan
# email: m7807031@gmail.com
# license: BSD-3

# OS import
import os
#from datetime import datetime, date
import time

# Copy module
import shutil

# Pandas import
import pandas as pd

# Office-docx lib import
import docx
from docx.shared import Cm, Pt  #add unit for word
#from docx.enum.text import WD_ALIGN_PARAGRAPH # deal with alignment
from docx.enum.table import WD_TABLE_ALIGNMENT

# ===== TODO ===== #
# 1. can't open large row excel
# 2. deal with folder's different name

# ===== Global Variables ===== #
input_folder = "input"
output_folder = "output"
row_switch = False
# ===== Global Function ===== #
def timeIt(func):
    def wrapper(*args, **kw):
        t1 = time.time()
        res = func(*args, **kw)
        t2 = time.time()
        print("** timeIt Report: {} took {:.2f}s".format(func.__name__, t2-t1))
    return wrapper

def ask_test_or_release():
    while 1:
        answer = input("(1)Test Mode or (2)Hard Mode:  ")
        try:
            answer = int(answer)
            if answer == 1:
                return False
            elif answer == 2:
                second_answer = input("WARNNING: Hard Mode will take lots of time, Sure? Y/N:  ")
                if second_answer == "Y" or second_answer == "y":
                    print("Have a good time :)")
                    return True
                else:
                    print("Input Error, Try again")
                    continue
        except:
            print("Input Error, Try again")
            continue

# ====== Functions ===== #
def output_folder_dir_list(input_folder_path):
    dir_list = os.listdir(input_folder_path)
    return dir_list


def xlsx2df(input_folder, tar_folder):
    tar_folder_list = os.listdir(input_folder+"/"+tar_folder)
    
    tar_xlsx = [item for item in tar_folder_list if item.endswith(u"_缺陷匯總表.xlsx")]


    raw_df = pd.read_excel(input_folder + "/" +tar_folder  + "/" + tar_xlsx[0])
    df = raw_df.loc[1::]
    df.columns = [u"序號",u"區間",u"距小號塔距離(m)",u"地物危險點座標", 
                  u"缺陷類型",u"缺陷級別",
                  u"實測平距",u"實測垂距",u"實測直線距離",
                  u"規範平距",u"規範垂距",u"規範直線距離",u"圖示"]
    
    df[u"區間"] = tar_folder
    #print(df.head(2))
    print(u"-> {} read.".format(tar_xlsx[0]) )
    return df


@timeIt
def df2docx(df, output_folder, tar_folder, row_switch:bool):
   
    docx_name = tar_folder + u"缺陷匯總表.docx"

    # get basic info from df
    if row_switch:
        wtf = df.shape[0]
    else:
        wtf = 100    

    row_length = wtf
    col_length = df.shape[1]

    print("-> Total Row: {}".format(row_length))
    

    # open temp docx file
    doc = docx.Document()

    # adjust border
    section_0 = doc.sections[0]
    section_0.left_margin = Cm(1.27)
    section_0.right_margin = Cm(1.27)
    section_0.top_margin = Cm(1.27)
    section_0.bottom_margin = Cm(1.27)


    table_0 = doc.add_table(rows=2 ,cols=col_length, style='Table Grid')
    table_0.alignment = WD_TABLE_ALIGNMENT.CENTER
    # deal with heading
    table_0.cell(0,0).merge(table_0.cell(1,0))
    table_0.cell(0,1).merge(table_0.cell(1,1))
    table_0.cell(0,2).merge(table_0.cell(1,2))
    table_0.cell(0,3).merge(table_0.cell(1,3))
    table_0.cell(0,4).merge(table_0.cell(1,4))
    table_0.cell(0,5).merge(table_0.cell(1,5))
    table_0.cell(0,6).merge(table_0.cell(0,8))
    table_0.cell(0,9).merge(table_0.cell(0,11))
    table_0.cell(0,12).merge(table_0.cell(1,12))

    table_0.cell(0,0).text = u"序號"
    table_0.cell(0,1).text = u"區間"
    table_0.cell(0,2).text = u"距小號塔距離(m)"
    table_0.cell(0,3).text = u"地物危險點座標"
    table_0.cell(0,4).text = u"缺陷類型"
    table_0.cell(0,5).text = u"缺陷級別"
    table_0.cell(0,6).text = u"實測距離(m)"
    table_0.cell(0,9).text = u"規範要求安全距離(m)"
    table_0.cell(0,12).text = u"圖示"

    table_0.cell(1,6).text = u"平距"
    table_0.cell(1,7).text = u"垂距"
    table_0.cell(1,8).text = u"直線距離"
    table_0.cell(1,9).text = u"平距"
    table_0.cell(1,10).text = u"垂距"
    table_0.cell(1,11).text = u"直線距離"
    
    print("---> table header ok.")

    
    table_1 = doc.add_table(rows=row_length ,cols=col_length, style='Table Grid')
    table_cells_1 = table_1._cells

    table_1.alignment = WD_TABLE_ALIGNMENT.CENTER
    print("---> table added.")


    for i in range(row_length):
        print("-> Row Processing: {} / {}". format(i+1, row_length))
        row_cells = table_cells_1[i*col_length : (i+1)*col_length]
        for j in range(len(row_cells)):
            insert_value = df.iloc[i, j]
            if j == 0:
                insert_value = str(int(insert_value))
            if j>=6 and j <=11:
                insert_value = "{:.3f}".format(insert_value)
            if str(insert_value) == "nan":
                insert_value = ""
            row_cells[j].text = str(insert_value)

        #Add text to row_cells

    doc.save(output_folder + "/" + tar_folder + "/" + docx_name)
    print("---> docx saved.")

# ====== main() ===== #
def main():
    dir_list = os.listdir(input_folder)
    for tar_folder in dir_list:
        
        # create output/sub-folder
        if not os.path.exists(output_folder + "/" + tar_folder):
            os.makedirs(output_folder + "/" + tar_folder)
        
        # rename file name
        print("------> {} is renaming...".format( input_folder + "/" + tar_folder + "/"))
        tar_dir_list = os.listdir(input_folder + "/" + tar_folder)
        for fn in tar_dir_list:
            old_name = input_folder + "/" + tar_folder + "/" + fn
            new_name = output_folder + "/" + tar_folder + "/" + tar_folder + fn.split("_")[-1]
            shutil.copy(old_name, new_name)
        print("------> {} renamed.".format( input_folder + "/" +tar_folder + "/"))


        # xlsx 2 docx
        print("------> {} is transfering to docx...".format( input_folder + "/" +tar_folder + "/"))
        single_df = xlsx2df(input_folder, tar_folder)
        df2docx(single_df, output_folder, tar_folder, row_switch=row_switch)
        print("------> {} transfered to docx.".format( input_folder + "/" +tar_folder + "/"))


if __name__=="__main__":
    print("***** START *****")

    row_switch = ask_test_or_release()
    main()

    print("***** DONE *****")
